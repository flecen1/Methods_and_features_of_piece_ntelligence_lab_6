
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from pathlib import Path


def set_paragraph_spacing(para, before=0, after=12, line_spacing=1.5):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = line_spacing


def add_title(doc, text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    set_paragraph_spacing(para, after=24)


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.name = "Times New Roman"
    set_paragraph_spacing(heading, before=12, after=6)


def add_paragraph(doc, text, bold=False, italic=False):
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        run.bold = bold
        run.italic = italic
    set_paragraph_spacing(para)


def add_list_item(doc, text, level=0):
    para = doc.add_paragraph(text, style=f'List {"Bullet" if level == 0 else "Bullet 2"}')
    for run in para.runs:
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
    set_paragraph_spacing(para, after=6)


def add_code_block(doc, code_text):
    para = doc.add_paragraph(code_text)
    para.paragraph_format.left_indent = Inches(0.5)
    for run in para.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(10)
    set_paragraph_spacing(para, before=6, after=6)


def add_image(doc, image_path, width=6.0):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if Path(image_path).exists():
        run = para.add_run()
        run.add_picture(str(image_path), width=Inches(width))
    else:
        run = para.add_run(f"[Рисунок: {image_path}]")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
    set_paragraph_spacing(para, before=12, after=12)


def main():
    doc = Document()

    # Title page
    add_title(doc, "МЕТОДИЧНІ ВКАЗІВКИ")
    add_title(doc, "до лабораторної роботи №6")
    add_title(doc, "з дисципліни «Методи та засоби штучного інтелекту»")
    add_paragraph(doc, "", bold=True)
    add_paragraph(doc, "Тема: Дискретна мережа Хопфілда", bold=True)
    add_paragraph(doc, "Мета: Ознайомитись із принципами роботи дискретної мережі Хопфілда, навчитись реалізовувати та досліджувати її для задач розпізнавання образів", bold=True)
    doc.add_page_break()

    # Table of contents placeholder
    add_heading(doc, "ЗМІСТ", level=1)
    add_list_item(doc, "1. Теоретичні відомості")
    add_list_item(doc, "2. Постановка задачі")
    add_list_item(doc, "3. Структура проєкту")
    add_list_item(doc, "4. Опис основних модулів")
    add_list_item(doc, "5. Порядок виконання роботи")
    add_list_item(doc, "6. Питання для самоконтролю")
    add_list_item(doc, "7. Література")
    doc.add_page_break()

    # 1. Теоретичні відомості
    add_heading(doc, "1. ТЕОРЕТИЧНІ ВІДОМОСТІ", level=1)
    add_paragraph(doc, "Мережа Хопфілда — це рекурентна нейронна мережа з повністю зв'язаними нейронами. Вона використовується для асоціативної пам'яті та розпізнавання образів.")
    add_paragraph(doc, "Основні характеристики мережі Хопфілда:", italic=True)
    add_list_item(doc, "Дискретні стани нейронів (+1, -1)")
    add_list_item(doc, "Повнозв'язна архітектура")
    add_list_item(doc, "Правило навчання Хебба")
    add_list_item(doc, "Асинхронне оновлення нейронів")
    add_list_item(doc, "Функція енергії (Ляпунова), що гарантує збіжність")

    add_heading(doc, "1.1 Математична основа", level=2)
    add_paragraph(doc, "Правило навчання Хебба з нормування:")
    add_code_block(doc, "W = (1/N) * sum_{k=1}^{P} x_k @ x_k^T")
    add_paragraph(doc, "де:")
    add_list_item(doc, "W — матриця ваг")
    add_list_item(doc, "N — кількість нейронів")
    add_list_item(doc, "P — кількість шаблонів")
    add_list_item(doc, "x_k — k-й біполярний шаблон (-1, +1)")

    add_paragraph(doc, "Асинхронне оновлення:")
    add_code_block(doc, "s_i(t+1) = sign( sum_j W_ij s_j(t) )")
    doc.add_page_break()

    # 2. Постановка задачі
    add_heading(doc, "2. ПОСТАНОВКА ЗАДАЧІ", level=1)
    add_paragraph(doc, "Необхідно реалізувати дискретну мережу Хопфілда для розпізнавання літер англійського алфавіту (розмір 7x5 пікселів).")
    add_paragraph(doc, "Завдання:")
    add_list_item(doc, "Навчити мережу на трьох літерах відповідно до варіанту")
    add_list_item(doc, "Дослідити стійкість до шуму")
    add_list_item(doc, "Дослідити роботу з «чужими» літерами (що не входили до навчальної вибірки)")
    add_list_item(doc, "Побудувати триптихи (еталон, зашумлений, відновлений)")
    add_list_item(doc, "Побудувати графіки зміни енергії")
    doc.add_page_break()

    # 3. Структура проєкту
    add_heading(doc, "3. СТРУКТУРА ПРОЄКТУ", level=1)
    add_paragraph(doc, "Проєкт має наступну структуру файлів:")
    add_list_item(doc, "hopfield.py — реалізація мережі Хопфілда")
    add_list_item(doc, "alphabet.py — завантаження шаблонів літер")
    add_list_item(doc, "variants.py — варіанти завдань")
    add_list_item(doc, "plotting.py — візуалізація результатів")
    add_list_item(doc, "main.py — консольний інтерфейс")
    add_list_item(doc, "gui.py — графічний інтерфейс")
    add_list_item(doc, "Alphabet.csv — файл з шаблонами літер")
    doc.add_page_break()

    # 4. Опис основних модулів
    add_heading(doc, "4. ОПИС ОСНОВНИХ МОДУЛІВ", level=1)

    add_heading(doc, "4.1 Модуль hopfield.py", level=2)
    add_paragraph(doc, "Містить основні функції для роботи з мережею:")
    add_list_item(doc, "train_weights(patterns_bipolar, zero_diagonal=True) — навчання мережі")
    add_list_item(doc, "recall_async(W, s0, max_sweeps, rng) — асинхронне відновлення образу")
    add_list_item(doc, "recall_with_energy_trace(W, s0, ...) — відновлення з логуванням енергії")
    add_list_item(doc, "energy(W, s) — розрахунок функції енергії")

    add_heading(doc, "4.2 Модуль alphabet.py", level=2)
    add_paragraph(doc, "Завантажує шаблони літер з файлу Alphabet.csv. Файл містить 35 рядків (7x5 пікселів) та 26 стовпців (літери A-Z).")

    add_heading(doc, "4.3 Модуль variants.py", level=2)
    add_paragraph(doc, "Містить варіанти завдань (20 варіантів, кожен з 3 літерами).")

    add_heading(doc, "4.4 Модуль main.py", level=2)
    add_paragraph(doc, "Консольний інтерфейс для запуску експериментів. Приклади використання:")
    add_code_block(doc, "python main.py --variant 1 --noise 0.4 --trials 50 --seed 42")
    add_code_block(doc, "python main.py --variant 1 --foreign V")
    add_code_block(doc, "python main.py --variant 1 --plot-triptych B --plot-file triptych_B.png")
    doc.add_page_break()

    # 5. Порядок виконання роботи
    add_heading(doc, "5. ПОРЯДОК ВИКОНАННЯ РОБОТИ", level=1)
    add_list_item(doc, "Перевірити наявність встановленого Python 3.8+")
    add_list_item(doc, "Встановити залежності: pip install -r requirements.txt")
    add_list_item(doc, "Визначити номер варіанту (за номером у списку групи або останніми цифрами залікової книжки)")
    add_list_item(doc, "Запустити шумовий експеримент: python main.py --variant N --noise 0.4 --trials 50")
    add_list_item(doc, "Запустити експеримент з «чужою» літерою: python main.py --variant N --foreign X")
    add_list_item(doc, "Побудувати триптихи для кожної навчальної літери")
    add_list_item(doc, "Побудувати графіки зміни енергії")
    add_list_item(doc, "Аналізуючи результати, оформити звіт відповідно до вимог")
    doc.add_page_break()

    # 6. Питання для самоконтролю
    add_heading(doc, "6. ПИТАННЯ ДЛЯ САМОКОНТРОЛЮ", level=1)
    add_list_item(doc, "Що таке мережа Хопфілда?")
    add_list_item(doc, "Які основні властивості мережі Хопфілда?")
    add_list_item(doc, "Що таке функція енергії в мережі Хопфілда?")
    add_list_item(doc, "Чим відрізняється синхронне та асинхронне оновлення нейронів?")
    add_list_item(doc, "Чому обнуляють діагональ матриці ваг?")
    add_list_item(doc, "Які обмеження має мережа Хопфілда?")
    add_list_item(doc, "Як оцінити якість роботи мережі з зашумленими даними?")
    doc.add_page_break()

    # 7. Література
    add_heading(doc, "7. ЛІТЕРАТУРА", level=1)
    add_list_item(doc, "1. Хайкін С. Нейронні мережі: повний курс. — М.: Вільямс, 2006.")
    add_list_item(doc, "2. Гудзенко В. М., Кривуля В. П. Штучний інтелект: навч. посібник. — К.: Вища школа, 2019.")
    add_list_item(doc, "3. Документація NumPy: https://numpy.org/doc/")
    add_list_item(doc, "4. Документація Matplotlib: https://matplotlib.org/")

    # Save document
    output_path = Path("МЕТОДИЧНІ_ВКАЗІВКИ_ЛР6_Хопфілд.docx")
    doc.save(output_path)
    print(f"Methodical documentation created: {output_path.absolute()}")


if __name__ == "__main__":
    main()
