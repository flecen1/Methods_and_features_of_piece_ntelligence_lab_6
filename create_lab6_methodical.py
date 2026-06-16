# -*- coding: utf-8 -*-
"""
Створення методичних вказівок для лабораторної роботи №6
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def add_heading(doc, text, level=1):
    """Додає заголовок"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.bold = True
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    return heading


def add_paragraph(doc, text, bold=False, italic=False):
    """Додає параграф"""
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
    return para


def add_code_block(doc, code):
    """Додає блок коду"""
    para = doc.add_paragraph(code)
    for run in para.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(10)
    para.paragraph_format.left_indent = Inches(0.5)
    return para


def add_image_placeholder(doc, description, width=6.0):
    """Додає плейсхолдер для зображення"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"[Рисунок: {description}]")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    return para


def create_methodical_doc():
    """Створює методичні вказівки"""
    doc = Document()
    
    # Титульна сторінка
    add_heading(doc, "МЕТОДИЧНІ ВКАЗІВКИ", level=1)
    add_heading(doc, "до виконання лабораторної роботи № 6", level=1)
    add_heading(doc, "«Дискретна мережа Хопфілда»", level=1)
    add_paragraph(doc, "")
    add_heading(doc, "Дисципліна «Методи та засоби штучного інтелекту»", level=2)
    add_paragraph(doc, "для студентів спеціальності 123 «Комп'ютерна інженерія»", bold=True)
    add_paragraph(doc, "")
    add_heading(doc, "ВСТУП", level=1)
    
    intro_text = """Метою лабораторної роботи є вивчення принципів функціонування дискретної мережі Хопфілда, навчання мережі на прикладі розпізнавання літер та дослідження її властивостей при роботі з зашумленими даними.

Мережа Хопфілда (Hopfield Network) — це рекурентна нейронна мережа з повною зв'язаністю нейронів, запропонована Джоном Хопфілдом у 1982 році. Мережа здатна зберігати та відновлювати образи (патерни), що робить її корисною для задач асоціативної пам'яті та розпізнавання образів.

Основні властивості мережі Хопфілда:
• Дискретні стани нейронів (-1, +1)
• Симетрична матриця ваг (W_ij = W_ji)
• Відсутність самозв'язків (діагональ матриці ваг дорівнює нулю)
• Асинхронне оновлення станів нейронів
• Наявність функції енергії (функції Ляпунова), яка гарантує збіжність"""
    
    add_paragraph(doc, intro_text)
    doc.add_page_break()
    
    # Теоретичні відомості
    add_heading(doc, "1. ТЕОРЕТИЧНІ ВІДОМОСТІ", level=1)
    
    add_heading(doc, "1.1. Структура мережі Хопфілда", level=2)
    
    theory_text1 = """Мережа Хопфілда складається з N нейронів, кожен з яких може перебувати в одному з двох станів: s_i ∈ {-1, +1}. Зв'язки між нейронами задаються матрицею ваг W розміром N×N.

Ключові особливості архітектури:
• Повна зв'язність: кожен нейрон з'єднаний з усіма іншими
• Симетрія: W_ij = W_ji
• Відсутність самозв'язків: W_ii = 0
• Біполярні стани: s_i ∈ {-1, +1}"""
    
    add_paragraph(doc, theory_text1)
    
    add_heading(doc, "1.2. Правило навчання Хебба", level=2)
    
    theory_text2 = """Навчання мережі Хопфілда відбувається за модифікованим правилом Хебба. Для набору з P патернів {x^1, x^2, ..., x^P}, де кожен патерн x^μ є N-вимірним біполярним вектором (x_i^μ ∈ {-1, +1}), елементи матриці ваг обчислюються як:

W_ij = (1/N) * Σ(μ=1 to P) x_i^μ * x_j^μ    для i ≠ j
W_ii = 0

де N — кількість нейронів (розмірність патерну).

Нормування на 1/N забезпечує стабільність процесу відновлення. Обнулення діагоналі запобігає самопідсиленню нейронів."""
    
    add_paragraph(doc, theory_text2)
    add_code_block(doc, "W = (1/N) * sum_{μ=1}^{P} x^μ @ (x^μ)^T\nnp.fill_diagonal(W, 0)")
    
    add_heading(doc, "1.3. Асинхронне оновлення станів", level=2)
    
    theory_text3 = """Процес відновлення (реколу) патерну в мережі Хопфілда відбувається асинхронно. На кожному кроці випадково обирається один нейрон i, і його стан оновлюється за правилом:

s_i(t+1) = sign(Σ(j=1 to N) W_ij * s_j(t))

де sign(x) = +1, якщо x ≥ 0, і -1, якщо x < 0.

Інші нейрони залишаються без змін: s_j(t+1) = s_j(t) для j ≠ i.

Асинхронність оновлення є ключовою для гарантії збіжності мережі до стабільного стану (атрактора)."""
    
    add_paragraph(doc, theory_text3)
    add_code_block(doc, "h = np.dot(W[i, :], s)\ns[i] = 1.0 if h >= 0.0 else -1.0")
    
    doc.add_page_break()
    
    # Продовження в наступному повідомленні...
    
    # Save document
    output_path = Path("МЕТОДИЧНІ_ВКАЗІВКИ_ЛР6_Хопфілд_ПОВНІ.docx")
    doc.save(output_path)
    print(f"Повні методичні вказівки створено: {output_path.absolute()}")


if __name__ == "__main__":
    main()
