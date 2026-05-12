# -*- coding: utf-8 -*-
"""
Генерує повний Word-файл «Методичні вказівки» до ЛР 6 (мережа Хопфілда)
за структурою методичних вказівок курсу M&SAI (розділи 1–13).

  pip install python-docx
  python build_methodical_docx.py

Вихід: МЕТОДИЧНІ_ВКАЗІВКИ_ЛР6_Хопфілд.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from variants import VARIANT_LETTERS

# Офіційний репозиторій практичної роботи 6 (M&SAI, мережа Хопфілда).
OFFICIAL_REPO_GIT = "https://github.com/flecen1/Methods_and_features_of_piece_ntelligence_lab_6.git"
OFFICIAL_REPO_WEB = "https://github.com/flecen1/Methods_and_features_of_piece_ntelligence_lab_6"
REPO_ROOT_FOLDER = "Methods_and_features_of_piece_ntelligence_lab_6"


def _center_run(p, text: str, *, bold: bool = False, size_pt: int = 14) -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size_pt)


def _h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str, *, bold: bool = False) -> None:
    r = doc.add_paragraph().add_run(text)
    r.bold = bold


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _mono(doc: Document, line: str) -> None:
    p = doc.add_paragraph(line)
    p.style = "Intense Quote"


def _task_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    tb = doc.add_table(rows=len(rows) + 1, cols=2)
    tb.style = "Table Grid"
    tb.rows[0].cells[0].text = "Завдання практичної роботи"
    tb.rows[0].cells[1].text = "Як виконати в проєкті (Python)"
    for i, (a, b) in enumerate(rows, start=1):
        tb.rows[i].cells[0].text = a
        tb.rows[i].cells[1].text = b


def _variant_table(doc: Document) -> None:
    tb = doc.add_table(rows=21, cols=4)
    tb.style = "Table Grid"
    for j, h in enumerate(("№ варіанту", "Літера 1", "Літера 2", "Літера 3")):
        tb.rows[0].cells[j].text = h
    for vid in range(1, 21):
        a, b, c = VARIANT_LETTERS[vid]
        row = tb.rows[vid]
        row.cells[0].text = str(vid)
        row.cells[1].text = a
        row.cells[2].text = b
        row.cells[3].text = c


def _files_appendix(doc: Document, rows: list[tuple[str, str]]) -> None:
    tb = doc.add_table(rows=len(rows) + 1, cols=2)
    tb.style = "Table Grid"
    tb.rows[0].cells[0].text = "Файл"
    tb.rows[0].cells[1].text = "Призначення"
    for i, (fn, desc) in enumerate(rows, start=1):
        tb.rows[i].cells[0].text = fn
        tb.rows[i].cells[1].text = desc


def build_document() -> Document:
    doc = Document()
    sty = doc.styles["Normal"]
    sty.font.name = "Times New Roman"
    sty.font.size = Pt(14)

    # --- Титульний блок ---
    for line, bold, sz in (
        ("МЕТОДИЧНІ ВКАЗІВКИ", True, 14),
        ("до практичної роботи з дисципліни", False, 14),
        ("«Моделі та системи штучного інтелекту»", True, 14),
        ("", False, 14),
        ("Практична робота", False, 14),
        (
            "«Дослідження дискретної мережі Хопфілда на задачі "
            "асоціативної пам’яті за допомогою засобів Python»",
            True,
            13,
        ),
    ):
        p = doc.add_paragraph()
        if line:
            _center_run(p, line, bold=bold, size_pt=sz)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    intro = (
        "Самостійна практична робота з курсу «Моделі та системи штучного інтелекту». "
        "Середовище виконання: мова Python (рекомендовано Python 3.10 або новіше).\n\n"
        "Усі студенти отримують матеріали з офіційного репозиторію на GitHub: "
        f"{OFFICIAL_REPO_WEB} — щоб версія коду збігалася у всіх і результати можна було відтворити."
    )
    ip = doc.add_paragraph(intro)
    ip.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # --- 1. Мета ---
    _h(doc, "1. Мета практичної роботи")
    _p(
        doc,
        "Ознайомитися з принципом роботи дискретної мережі Хопфілда як асоціативної пам’яті; "
        "набути практичних навичок навчання мережі за правилом Хебба, запуску сценаріїв у Python, "
        "аналізу відновлення спотворених образів (літер), побудови ілюстрацій і роботи з "
        "графічним інтерфейсом програми.",
    )

    # --- 2. Зміст ---
    _h(doc, "2. Зміст практичної роботи: завдання та реалізація у проєкті")
    _p(
        doc,
        "Робота складається з чотирьох основних завдань і таблиці з двадцятьма варіантами "
        "наборів літер для навчання (розділ 7). Матеріали потрібно брати лише з офіційного "
        f"репозиторію ({OFFICIAL_REPO_WEB}). Нижче — що зробити і який файл проєкту для цього використати.",
    )
    _task_table(
        doc,
        [
            (
                "1. Демонстраційний приклад мережі Хопфілда (мала розмірність).",
                "Файл main.py з параметром --demo-toy. Перевіряється стабільність еталонних "
                "станів; деталі реалізації — у hopfield.py (функція demo_toy).",
            ),
            (
                "2. Навчити мережу на трьох літерах згідно з варіантом (табл. розділу 7); "
                "дослідити відновлення після шуму.",
                "main.py: параметри --variant 1…20, --noise σ, --trials K, опційно --seed. "
                "Літери та правила призначення варіанта — variants.py; дані літер — Alphabet.csv, "
                "модуль alphabet.py.",
            ),
            (
                "3. Подати на вхід літеру, якої не було серед трьох навчених («чужа» літера).",
                "main.py: --variant N та --foreign L (літера A…Z поза навчальною трійкою). "
                "У звіті проаналізувати, до якого з еталонів «тягне» мережа (поля JSON або текст консолі).",
            ),
            (
                "4. Побудувати рисунки для звіту та/або виконати експерименти з графічного інтерфейсу.",
                "main.py: --plot-triptych, --plot-file, --plot-energy, --no-show (модуль plotting.py). "
                "Або gui.py — вікно з полями параметрів і кнопками запуску тих самих сценаріїв.",
            ),
        ],
    )
    doc.add_paragraph()

    # --- 3. Формулювання ---
    _h(doc, "3. Формулювання завдань практичної роботи")
    _p(
        doc,
        "Тема: дослідження дискретної мережі Хопфілда на задачі асоціативної пам’яті "
        "(розпізнавання спотворених бінарних образів літер) засобами Python.",
        bold=True,
    )
    _p(doc, "Завдання:", bold=True)
    for i, t in enumerate(
        (
            "Виконати демонстраційний сценарій --demo-toy і зафіксувати результат у звіті.",
            "Навчити мережу на трьох літерах свого варіанту; провести серії експериментів "
            "із різним рівнем шуму σ; оцінити частку успішних відновлень (accuracy) для кожної літери.",
            "Подати «чужу» літеру (не з навчальної трійки); описати кінцевий стан і найближчий "
            "навчений еталон за відстанню Геммінга.",
            "Побудувати хоча б один рисунок (еталон / зашумлений вхід / результат після релаксації) "
            "або скріншот gui.py; за бажанням — графік енергії при релаксації.",
        ),
        start=1,
    ):
        doc.add_paragraph(f"{i}.\t{t}", style="List Number")

    _p(
        doc,
        "У таблиці варіантів (розділ 7) номер рядка 1…20 відповідає параметру --variant. "
        "Правила призначення варіанта студенту (список групи, залікова книжка тощо) визначає викладач; "
        "у коді наведено приклади функцій variant_from_journal_order та variant_from_gradebook_last_two_digits.",
    )

    # --- 4. Покрокова інструкція ---
    _h(doc, "4. Покрокова інструкція для студента (простими словами)")
    _p(
        doc,
        "Виконуйте кроки по порядку. Після кожного кроку заносьте у звіт те, що вимагає викладач: "
        "числа, скріншот або короткий коментар.",
    )

    _p(doc, "Крок 0. Отримати проєкт і підготувати середовище", bold=True)
    _bullet(
        doc,
        "Установіть Python (3.10+) та Git (за бажанням). Завантажте код лише з офіційного репозиторію "
        f"на GitHub ({OFFICIAL_REPO_WEB}) — клонуванням або кнопкою Code → Download ZIP.",
    )
    _p(doc, "Приклад клонування в PowerShell (папка на робочому столі):", bold=False)
    _mono(doc, "cd $env:USERPROFILE\\Desktop")
    _mono(doc, f"git clone {OFFICIAL_REPO_GIT}")
    _mono(doc, f"cd {REPO_ROOT_FOLDER}")
    _p(
        doc,
        "Якщо проєкт отримано як архів — розпакуйте його та перейдіть у корінь папки з файлами "
        "main.py, hopfield.py тощо.",
    )
    _p(doc, "Установка залежностей (див. також розділ 8):", bold=True)
    for ln in (
        "python -m venv .venv",
        ".\\.venv\\Scripts\\Activate.ps1",
        "pip install -r requirements.txt",
        "$env:PYTHONUTF8=1",
    ):
        _mono(doc, ln)
    _p(
        doc,
        "Переконайтеся, що команда python main.py --help виводить довідку без помилки. "
        "Файл Alphabet.csv має перебувати в корені проєкту (або вкажіть шлях параметром --alphabet).",
    )

    _p(doc, "Крок 1 — завдання 1 (демонстраційний приклад)", bold=True)
    _mono(doc, "python main.py --demo-toy --json")
    _p(
        doc,
        "У звіті: матриця ваг (поле weights) або скрін консолі, коротко — чи всі чотири еталонні "
        "образи залишаються стабільними (matches_target). Зверніть увагу: для цього малого прикладу "
        "у програмі діагональ матриці ваг спеціально не обнуляється (інакше W стала б нульовою); "
        "для літер (35 нейронів) обнулення діагоналі застосовується завжди.",
    )

    _p(doc, "Крок 2 — завдання 2 (шум і точність для свого варіанту)", bold=True)
    _p(doc, "Дізнайтеся свій номер варіанту 1…20 (розділ 7, журнал групи). Підставте N замість номера:", bold=False)
    _mono(doc, "python main.py --variant N --noise 0.35 --trials 50 --seed 7 --json")
    _p(
        doc,
        "Повторіть для кількох значень σ (наприклад 0.2; 0.4; 0.6). У звіті — таблиця: стовпці "
        "σ, літера, accuracy, за бажанням converged_fraction. Опишіть, як зростання шуму впливає "
        "на відновлення.",
    )

    _p(doc, "Крок 3 — завдання 3 (чужа літера)", bold=True)
    _p(
        doc,
        "Оберіть літеру A…Z, якої немає серед трьох літер вашого варіанту (див. variants.py або табл. розділу 7). "
        "Приклад (варіант 1: B, D, Y — можна взяти V):",
        bold=False,
    )
    _mono(doc, "python main.py --variant 1 --foreign V --json")
    _p(
        doc,
        "У звіті: поля closest_trained_letter, hamming_bipolar_to_trained, короткий висновок.",
    )

    _p(doc, "Крок 4 — завдання 4 (рисунок і/або GUI)", bold=True)
    _mono(doc, "mkdir figures")
    _mono(
        doc,
        "python main.py --variant N --plot-triptych B --noise 0.55 "
        "--plot-file figures\\variantN_B.png --plot-energy --no-show",
    )
    _p(
        doc,
        "Літера після --plot-triptych має входити до трійки вашого варіанту. Для графічного інтерфейсу:",
        bold=False,
    )
    _mono(doc, "python gui.py")
    _p(doc, "Зробіть скріншот вікна з результатом для розділу «хід роботи».", bold=False)

    _p(doc, "Крок 5. Оформлення звіту", bold=True)
    _p(doc, "Зберіть скріншоти, таблиці експериментів і рисунки. Заповніть розділи згідно з розділом 10 цих вказівок.")

    # --- 5. Репозиторій ---
    _h(doc, "5. З чого складається репозиторій (основні файли)")
    for line in (
        "main.py — консольний запуск усіх сценаріїв;",
        "hopfield.py — навчання ваг (Хебб), асинхронна релаксація, шумний старт, енергія;",
        "alphabet.py — читання Alphabet.csv і формування вектора літери;",
        "variants.py — таблиця варіантів 1…20 (три літери) і допоміжні функції призначення варіанта;",
        "plotting.py — відображення 7×5, triptych, графік енергії;",
        "gui.py — графічний інтерфейс;",
        "Alphabet.csv — бінарні шаблони літер (35×26);",
        "requirements.txt — залежності (numpy, matplotlib, за потреби python-docx для генерації вказівок).",
    ):
        _bullet(doc, line)

    # --- 6. Теорія ---
    _h(doc, "6. Теоретичні відомості (стисло)")
    _p(
        doc,
        "Мережа Хопфілда (дискретна, повнозв’язна) зберігає набір біполярних образів як стійкі стани "
        "динаміки. Навчання за правилом Хебба: ваги залежать від кореляцій між компонентами образів; "
        "для стабільності та уникнення самозв’язків діагональ матриці ваг часто обнуляється після обчислення W. "
        "Оновлення нейронів асинхронне: за один мікрокрок змінюється один компонент вектора стану за знаком "
        "локального поля. Існує функція енергії (Ляпунова); при асинхронних оновленнях вона не зростає — "
        "траєкторія сходиться до атрактора (навчений еталон, змішаний стан або «хибний» атрактор). "
        "Ємність пам’яті обмежена; схожі образи й сильний шум підвищують ризик помилок відновлення.",
    )

    # --- 7. Варіанти ---
    _h(doc, "7. Таблиця варіантів (три літери для навчання, відповідає --variant 1…20)")
    _variant_table(doc)
    doc.add_paragraph()
    _p(doc, "Призначення варіанта студенту (приклади правил):", bold=True)
    _bullet(
        doc,
        "А: номер у алфавітному списку групи n = 1, 2, … → варіант v = ((n − 1) mod 20) + 1.",
    )
    _bullet(
        doc,
        "Б: останні дві цифри залікової книжки як число k (0…99); у програмі: k mod 100; якщо k = 0 → варіант 20; "
        "інакше v = ((k − 1) mod 20) + 1.",
    )
    _bullet(doc, "В: номер варіанта задає викладач у журналі / LMS.")

    # --- 8. Установка ---
    _h(doc, "8. Отримання проєкту та установка бібліотек")
    _p(
        doc,
        f"Адреса репозиторію: {OFFICIAL_REPO_WEB}. Після клонування або розпакування ZIP перейдіть у корінь "
        f"папки {REPO_ROOT_FOLDER} (там, де лежить main.py) і виконайте команди (PowerShell):",
    )
    _mono(doc, f"cd $env:USERPROFILE\\Desktop\\{REPO_ROOT_FOLDER}")
    for ln in (
        "python -m venv .venv",
        ".\\.venv\\Scripts\\Activate.ps1",
        "pip install -r requirements.txt",
    ):
        _mono(doc, ln)
    _mono(doc, "$env:PYTHONUTF8=1")
    _p(
        doc,
        "Якщо потрібно відновити стандартний Alphabet.csv з шрифту Adafruit 5×7: "
        "python build_alphabet_from_glcdfont.py (може знадобитися доступ до Інтернету для завантаження glcdfont.c).",
    )

    # --- 9. Довідник команд ---
    _h(doc, "9. Довідник команд (коротко)")
    _mono(
        doc,
        "python main.py --demo-toy [--json]",
    )
    _mono(
        doc,
        "python main.py --variant N [--noise σ] [--trials K] [--seed S] [--alphabet шлях\\Alphabet.csv] [--json]",
    )
    _mono(
        doc,
        "python main.py --variant N --foreign L [--seed S] [--alphabet …] [--json]",
    )
    _mono(
        doc,
        "python main.py --variant N --plot-triptych L --noise σ --plot-file шлях.png "
        "[--plot-energy] [--no-show]",
    )
    _mono(doc, "python gui.py")

    # --- 10. Звіт ---
    _h(doc, "10. Зміст звіту студента")
    _p(doc, "У звіті обов’язково мають бути такі частини (у зручному для вас порядку):", bold=False)
    parts = (
        "Титульна сторінка.",
        "Мета роботи.",
        "Завдання (формулювання з цих вказівок або з бланку викладача).",
        "Опис проєкту: посилання на репозиторій / джерело коду, перелік основних файлів (див. розділ 5 і додаток 12).",
        "Хід роботи зі скріншотами та таблицями: демо --demo-toy; експерименти з шумом; чужа літера; рисунки або GUI.",
        "Висновки (що спостерігалося при зміні σ, чи збігається «чужа» літера з найближчим еталоном тощо).",
    )
    for i, t in enumerate(parts, start=1):
        doc.add_paragraph(f"{i}.\t{t}", style="List Number")
    _p(
        doc,
        "У пункті «хід роботи» логічно розташувати: вивід після кроку 1, таблиці після кроку 2, "
        "аналіз --foreign після кроку 3, рисунок або скрін gui.py після кроку 4.",
    )

    # --- 11. Контрольні питання ---
    _h(doc, "11. Контрольні питання та зразкові відповіді")
    qa = [
        (
            "1. Які задачі розв’язує мережа Хопфілда?",
            "Асоціативна пам’ять, відновлення спотворених дискретних образів; динаміка веде стан "
            "до локальних мінімумів енергії (атракторів).",
        ),
        (
            "2. Які умови забезпечують стійкість класичної моделі (стисло)?",
            "Симетрична матриця ваг, нуль на діагоналі, асинхронне оновлення; енергія не зростає, "
            "збіжність до стійкого стану.",
        ),
        (
            "3. Назвіть основні недоліки мережі Хопфілда як пам’яті.",
            "Обмежена ємність, хибні атрактори, чутливість до сильного шуму та до візуально схожих образів.",
        ),
        (
            "4. Що таке енергетична функція в цьому контексті?",
            "Скалярна величина, яка характеризує «якість» поточного стану; при асинхронних кроках оновлення "
            "енергія не збільшується.",
        ),
        (
            "5. За яких умов зупиняється ітераційний процес у програмі?",
            "Після повного асинхронного проходу по нейронах стан не змінюється, або досягнуто внутрішнього "
            "ліміту кількості проходів (max_sweeps).",
        ),
    ]
    for q, a in qa:
        _p(doc, q, bold=True)
        _p(doc, f"Відповідь: {a}")
        doc.add_paragraph()

    # --- 12. Додаток ---
    _h(doc, "12. Додаток: відповідність файлів проєкту")
    _files_appendix(
        doc,
        [
            ("variants.py", "варіанти 1…20 (три літери), функції призначення варіанта"),
            ("hopfield.py", "навчання ваг, релаксація, шум, енергія"),
            ("alphabet.py", "завантаження шаблонів літер"),
            ("plotting.py", "візуалізація 7×5, triptych, енергія"),
            ("main.py", "консольний запуск"),
            ("gui.py", "графічний інтерфейс"),
            ("Alphabet.csv", "дані літер"),
            ("requirements.txt", "залежності"),
        ],
    )

    # --- 13. Література ---
    _h(doc, "13. Рекомендована література")
    for ref in (
        "Hopfield J. J. Neural networks and physical systems with emergent collective computational abilities. "
        "PNAS, 1982.",
        "Рутковська Д., Пилинський М., Рутковський Л. Нейронні мережі, генетичні алгоритми та нечіткі системи. "
        "— М.: Гаряча лінія — Телеком, 2006.",
        "Кононюк А. Ю. Нейронні мережі і генетичні алгоритми. — К.: Корнійчук, 2008.",
        "Документація Python: https://docs.python.org/3/",
        "Документація NumPy: https://numpy.org/doc/stable/",
    ):
        _bullet(doc, ref)

    return doc


def main() -> None:
    out = Path(__file__).resolve().parent / "МЕТОДИЧНІ_ВКАЗІВКИ_ЛР6_Хопфілд.docx"
    doc = build_document()
    doc.save(out)
    print("Записано:", out)


if __name__ == "__main__":
    main()
