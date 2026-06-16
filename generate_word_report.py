
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def add_heading(doc, text, level=1, bold=True, size=16, centered=False):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    run = heading.runs[0]
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return heading


def add_paragraph(doc, text, size=12, bold=False, italic=False):
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"
        run.bold = bold
        run.italic = italic
    return para


def add_table(doc, headers, data, width=1.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
        hdr_cells[i].paragraphs[0].runs[0].font.name = "Times New Roman"
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
            row_cells[i].paragraphs[0].runs[0].font.name = "Times New Roman"
    return table


def add_image(doc, image_path, width=6.0):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if Path(image_path).exists():
        run = para.add_run()
        run.add_picture(str(image_path), width=Inches(width))
    else:
        run = para.add_run(f"[Місце для рисунка: {image_path}]")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(128, 128, 128)


def main():
    with open("word_report_log.txt", "w", encoding="utf-8") as log_file:
        def log(message):
            print(message)
            log_file.write(message + "\n")

        try:
            log("Starting Word document generation...")
            doc = Document()
            log("Document object created")

            # Title page
            add_heading(doc, "ЗВІТ", level=1, size=18, centered=True)
            add_heading(doc, "до лабораторної роботи №6", level=2, size=16, centered=True)
            add_heading(doc, "з дисципліни «Методи та засоби штучного інтелекту»", level=2, size=14, centered=True)
            doc.add_paragraph()
            doc.add_paragraph("---")
            doc.add_paragraph()

            add_paragraph(doc, "Тема: Мережа Хопфілда", bold=True)
            add_paragraph(doc, "Виконав: студент групи КНД-31")
            add_paragraph(doc, "Петренко Ярослав Олексійович")
            add_paragraph(doc, "(номер у списку групи: 1)")
            add_paragraph(doc, "Номер варіанту: 1")
            add_paragraph(doc, "Очікувана оцінка: 5")
            doc.add_paragraph()
            doc.add_paragraph("---")
            doc.add_paragraph()

            add_heading(doc, "Місто Київ — 2025 рік", level=3, size=12, centered=True)

            # New page
            doc.add_page_break()
            log("Title page added")

            # Content
            add_heading(doc, "Мета", level=1)
            add_paragraph(doc, "Навчитись працювати з дискретною мережею Хопфілда, навчити мережу на трьох літерах, дослідити її роботу на зашумлених входах та з «чужими» літерами.")
            log("Meta section added")

            add_heading(doc, "Теоретичні відомості", level=1)
            add_paragraph(doc, "Дискретна мережа Хопфілда — це рекурентна нейронна мережа, що складається з N нейронів. Стан кожного нейрона може бути біполярним (+1 або -1). Мережа здатна зберігати декілька образів (патернів) як атрактори, а потім відновлювати оригінальний образ із зашумленого входу.")

            add_heading(doc, "Формули", level=2)
            add_paragraph(doc, "1. Біполярне представлення:", italic=True)
            add_paragraph(doc, "   x = 2b - 1")
            add_paragraph(doc, "2. Навчання за правилом Хебба (нормоване):", italic=True)
            add_paragraph(doc, "   W = (1/N) * sum_{k=1}^{P} x_k x_k^T")
            add_paragraph(doc, "   де N — кількість нейронів, P — кількість образів, x_k — k-ий біполярний образ.")
            add_paragraph(doc, "3. Асинхронне оновлення:", italic=True)
            add_paragraph(doc, "   s_i(t+1) = +1, якщо sum_j W_ij s_j(t) ≥ 0, інакше -1")
            add_paragraph(doc, "4. Функція енергії (функція Ляпунова):", italic=True)
            add_paragraph(doc, "   E(s) = -0.5 * s^T W s")
            log("Theory section added")

            add_heading(doc, "Виконання роботи", level=1)
            add_heading(doc, "Інструменти", level=2)
            add_paragraph(doc, "• Мова програмування: Python 3.11")
            add_paragraph(doc, "• Бібліотеки: NumPy, Matplotlib, Tkinter, python-docx")

            add_heading(doc, "Основні результати", level=1)
            add_heading(doc, "1. Варіант та навчальні літери", level=2)
            add_paragraph(doc, "Мій варіант — 1, навчальні літери: B, D, Y.")

            add_heading(doc, "2. Шумовий експеримент", level=2)
            add_paragraph(doc, "Проведено експеримент з силою шуму σ = 0.4, 50 випробувань на літеру:")
            noise_data = [
                ["B", "0.00", "1.00"],
                ["D", "0.18", "1.00"],
                ["Y", "0.66", "1.00"],
            ]
            add_table(doc, ["Літера", "Точність (Accuracy)", "Частота збіжності"], noise_data)

            add_paragraph(doc, "Висновки: Літера Y має найкращу точність (66%), літера B — найгіршу (0%). Це можна пояснити тим, що літери B і D мають схожі шаблони, тому мережа частіше їх плутає. У всіх випадках мережа збігалася до стабільного стану.")
            log("Noise experiment section added")

            add_heading(doc, "3. Триптихи", level=2)
            add_paragraph(doc, "Триптихи показують три стани: еталон, зашумлений вхід та відновлений образ:")
            add_image(doc, "figures/triptych_B.png")
            add_paragraph(doc, "Триптых для літери B")
            add_image(doc, "figures/triptych_D.png")
            add_paragraph(doc, "Триптых для літери D")
            add_image(doc, "figures/triptych_Y.png")
            add_paragraph(doc, "Триптых для літери Y")

            add_heading(doc, "4. Графіки енергії", level=2)
            add_paragraph(doc, "Графіки показують зміну енергії мережі під час асинхронної релаксації:")
            add_image(doc, "figures/energy_B.png")
            add_paragraph(doc, "Графік енергії для літери B")
            add_image(doc, "figures/energy_D.png")
            add_paragraph(doc, "Графік енергії для літери D")
            add_image(doc, "figures/energy_Y.png")
            add_paragraph(doc, "Графік енергії для літери Y")
            log("Triptychs and energy section added")

            add_heading(doc, "5. «Чужа» літера", level=2)
            add_paragraph(doc, "Використовували літеру V (не входила до навчальної вибірки):")
            add_paragraph(doc, "• Обучені літери: B, D, Y")
            add_paragraph(doc, "• Чужа літера: V")
            add_paragraph(doc, "• Найближчий еталон після релаксації: D")
            add_paragraph(doc, "• Кількість проходів до збіжності: 2")
            add_paragraph(doc, "• Збіжність: так")
            log("Foreign letter section added")

            add_heading(doc, "Висновки", level=1)
            add_paragraph(doc, "1. Успішно реалізовано дискретну мережу Хопфілда з асинхронним оновленням.")
            add_paragraph(doc, "2. Мережа здатна відновлювати літери з зашумлених входів, але точність залежить від схожості навчальних образів.")
            add_paragraph(doc, "3. Функція енергії монотонно спадає, що підтверджує теоретичні положення про збіжність мережі.")
            add_paragraph(doc, "4. «Чужі» літери (не з навчальної вибірки) мають тенденцію сходитися до найближчого навчального еталону (за відстанню Хеммінга).")
            add_paragraph(doc, "5. Мережа Хопфілда може бути використана для задач асоціативної пам'яті та розпізнавання простих образів.")
            log("Conclusions section added")

            # Save document
            output_path = Path("ЗВІТ_ЛР6_Петренко_Ярослав_КНД-31.docx")
            doc.save(output_path)
            log(f"Word document created: {output_path.absolute()}")
        except Exception as e:
            log(f"ERROR: {str(e)}")
            import traceback
            log(traceback.format_exc())

if __name__ == "__main__":
    main()
