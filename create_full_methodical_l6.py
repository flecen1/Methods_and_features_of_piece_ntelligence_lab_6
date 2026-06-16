# -*- coding: utf-8 -*-
"""
Створення повних методичних вказівок для лабораторної роботи №6
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


def set_cell_border(cell, **kwargs):
    """Встановлює границі для комірки таблиці"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    


def create_document():
    """Створює документ методичних вказівок"""
    doc = Document()
    
    # Налаштування сторінки (A4)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    
    return doc


def add_title_block(doc):
    """Додає титульний блок"""
    # Міністерство
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True
    
    # Університет
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ДЕРЖАВНИЙ УНІВЕРСИТЕТ ІНФРАСТРУКТУРИ ТА ТЕХНОЛОГІЙ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True
    
    # Факультет
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Факультет комп'ютерних технологій та систем")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True
    
    # Кафедра
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Кафедра комп'ютерних та інформаційних технологій")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True
    
    # Порожній рядок
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Затверджую
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("ЗАТВЕРДЖУЮ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True
    
    # Завідувач кафедри
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Завідувач кафедри КІТ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    
    # Підпис
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("_______________")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    
    # Дата
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("«___»_____________2025 р.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    
    # Порожній рядок
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Методичні вказівки
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("МЕТОДИЧНІ ВКАЗІВКИ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True
    
    # до виконання лабораторної роботи
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("до виконання лабораторної роботи")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True
    
    # № 6
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("№ 6")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True
    
    # Дискретна мережа Хопфілда
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("«Дискретна мережа Хопфілда»")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True
    run.italic = True
    
    # Порожній рядок
    doc.add_paragraph()
    
    # Дисципліна
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Дисципліна «Методи та засоби штучного інтелекту»")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True
    
    # Спеціальність
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("для здобувачів ступеня бакалавра за спеціальністю")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    
    # 123 КІ
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("123 «Комп'ютерна інженерія»")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True
    
    # Порожній рядок
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Київ
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Київ — 2025")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True


def add_page_with_content(doc):
    """Додає сторінку з основним змістом"""
    # Розрив сторінки
    doc.add_page_break()
    
    # Мета та завдання
    add_heading(doc, "МЕТА ТА ЗАВДАННЯ", level=1)
    
    add_heading(doc, "Мета роботи:", level=2)
    add_paragraph(doc, "Вивчити принципи функціонування дискретної мережі Хопфілда; навчитися реалізовувати та досліджувати мережу для задач розпізнавання образів.")
    
    add_heading(doc, "Завдання:", level=2)
    tasks = [
        "1. Ознайомитися з теоретичними основами мережі Хопфілда.",
        "2. Реалізувати дискретну мережу Хопфілда з асинхронним оновленням.",
        "3. Навчити мережу на трьох літерах відповідно до варіанту.",
        "4. Дослідити роботу мережі з зашумленими входами.",
        "5. Дослідити роботу мережі з «чужими» літерами.",
        "6. Побудувати графіки зміни енергії мережі.",
        "7. Оформити звіт відповідно до вимог."
    ]
    for task in tasks:
        add_paragraph(doc, task)
    
    # Теоретичні відомості
    doc.add_page_break()
    add_heading(doc, "ТЕОРЕТИЧНІ ВІДОМОСТІ", level=1)
    
    add_heading(doc, "1.1. Мережа Хопфілда", level=2)
    theory1 = """Мережа Хопфілда (Hopfield Network) — це рекурентна нейронна мережа, запропонована Джоном Хопфілдом у 1982 році. Вона є однією з найважливіших архітектур нейронних мереж для асоціативної пам'яті.

Основні характеристики мережі Хопфілда:
• Дискретні стани нейронів (-1, +1)
• Симетрична матриця ваг (W_ij = W_ji)
• Відсутність самозв'язків (W_ii = 0)
• Асинхронне оновлення нейронів
• Гарантована збіжність до стабільного стану"""
    add_paragraph(doc, theory1)
    
    add_heading(doc, "1.2. Математична модель", level=2)
    theory2 = """Мережа складається з N нейронів, кожен з яких може перебувати в одному з двох станів: s_i ∈ {-1, +1}. Стан мережі визначається вектором s = (s_1, s_2, ..., s_N).

Зв'язки між нейронами задаються матрицею ваг W розміром N×N. Елемент W_ij визначає силу зв'язку від нейрона j до нейрона i.

Ключова властивість мережі Хопфілда — симетрія матриці ваг:
W_ij = W_ji

Також передбачається відсутність самозв'язків:
W_ii = 0"""
    add_paragraph(doc, theory2)
    
    add_heading(doc, "1.3. Правило навчання Хебба", level=2)
    theory3 = """Навчання мережі Хопфілда відбувається за модифікованим правилом Хебба. Для набору з P патернів {x^1, x^2, ..., x^P}, де кожен патерн x^μ є N-вимірним біполярним вектором (x_i^μ ∈ {-1, +1}), елементи матриці ваг обчислюються як:

W_ij = (1/N) * Σ(μ=1 to P) x_i^μ * x_j^μ    для i ≠ j
W_ii = 0

де N — кількість нейронів (розмірність патерну).

Нормування на 1/N забезпечує стабільність процесу відновлення. Обнулення діагоналі запобігає самопідсиленню нейронів."""
    add_paragraph(doc, theory3)
    add_code_block(doc, "# Реалізація на Python\nimport numpy as np\n\ndef train_weights(patterns_bipolar):\n    N, P = patterns_bipolar.shape\n    W = np.einsum('ik,jk->ij', patterns_bipolar, patterns_bipolar) / N\n    np.fill_diagonal(W, 0.0)\n    return W")
    
    add_heading(doc, "1.4. Асинхронне оновлення станів", level=2)
    theory4 = """Процес відновлення (реколу) патерну в мережі Хопфілда відбувається асинхронно. На кожному кроці випадково обирається один нейрон i, і його стан оновлюється за правилом:

s_i(t+1) = sign(Σ(j=1 to N) W_ij * s_j(t))

де sign(x) = +1, якщо x ≥ 0, і -1, якщо x < 0.

Інші нейрони залишаються без змін: s_j(t+1) = s_j(t) для j ≠ i.

Асинхронність оновлення є ключовою для гарантії збіжності мережі до стабільного стану (атрактора)."""
    add_paragraph(doc, theory4)
    add_code_block(doc, "# Асинхронне оновлення\ndef recall_async(W, s0, max_sweeps=2000):\n    s = s0.copy()\n    n = len(s)\n    rng = np.random.default_rng()\n    \n    for sweep in range(max_sweeps):\n        prev = s.copy()\n        for i in rng.permutation(n):\n            h = np.dot(W[i, :], s)\n            s[i] = 1.0 if h >= 0.0 else -1.0\n        if np.array_equal(s, prev):\n            break\n    return s")
    
    add_heading(doc, "1.5. Функція енергії Ляпунова", level=2)
    theory5 = """Для аналізу динаміки мережі Хопфілда використовується функція енергії (функція Ляпунова), яка визначається як:

E(s) = -0.5 * Σ(i=1 to N) Σ(j=1 to N) W_ij * s_i * s_j

або у матричній формі:

E(s) = -0.5 * s^T * W * s

Ключова властивість: при асинхронному оновленні енергія мережі не зростає (ΔE ≤ 0), що гарантує збіжність до локального мінімуму (стабільного стану)."""
    add_paragraph(doc, theory5)
    add_code_block(doc, "# Розрахунок енергії\ndef energy(W, s):\n    return -0.5 * np.dot(s, np.dot(W, s))")
    
    doc.add_page_break()
    
    # Решта змісту...
    
    return doc


def main():
    """Головна функція"""
    print("Створення методичних вказівок для лабораторної роботи №6...")
    
    doc = create_document()
    add_title_block(doc)
    doc = add_content(doc)
    
    # Збереження
    output_path = Path("МЕТОДИЧНІ_ВКАЗІВКИ_ЛР6_Хопфілд_ПОВНІ.docx")
    doc.save(output_path)
    
    print(f"\nМетодичні вказівки збережено: {output_path.absolute()}")
    print("Готово!")


if __name__ == "__main__":
    main()
