# -*- coding: utf-8 -*-
"""
Створення звіту студента у форматі DOCX для лабораторної роботи №6
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def create_student_report():
    """Створює звіт студента"""
    doc = Document()
    
    # Титульна сторінка
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ЗВІТ")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("до лабораторної роботи №6")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("з дисципліни «Методи та засоби штучного інтелекту»")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Тема
    p = doc.add_paragraph()
    run = p.add_run("Тема: Дискретна мережа Хопфілда")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    
    doc.add_paragraph()
    
    # Виконав
    p = doc.add_paragraph()
    run = p.add_run("Виконав: студент групи КНД-31")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    
    p = doc.add_paragraph()
    run = p.add_run("Петренко Ярослав Олексійович")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    
    p = doc.add_paragraph()
    run = p.add_run("(номер у списку групи: 1)")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("Номер варіанту: 1")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    
    p = doc.add_paragraph()
    run = p.add_run("Очікувана оцінка: 5")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Місто Київ
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Місто Київ — 2025 рік")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    
    doc.add_page_break()
    
    # Мета
    add_heading(doc, "МЕТА", level=1)
    add_paragraph(doc, "Навчитись працювати з дискретною мережею Хопфілда, навчити мережу на трьох літерах, дослідити її роботу на зашумлених входах та з «чужими» літерами.")
    
    # Теоретичні відомості
    add_heading(doc, "ТЕОРЕТИЧНІ ВІДОМОСТІ", level=1)
    
    add_paragraph(doc, "Дискретна мережа Хопфілда — це рекурентна нейронна мережа, що складається з N нейронів. Стан кожного нейрона може бути біполярним (+1 або -1). Мережа здатна зберігати декілька образів (патернів) як атрактори, а потім відновлювати оригінальний образ із зашумленого входу.")
    
    add_heading(doc, "Формули", level=2)
    
    add_paragraph(doc, "1. Біполярне представлення:", italic=True)
    add_paragraph(doc, "   x = 2b - 1")
    
    add_paragraph(doc, "2. Навчання за правилом Хебба (нормоване):", italic=True)
    add_paragraph(doc, "   W = (1/N) * Σ(k=1 to P) x_k @ x_k^T")
    add_paragraph(doc, "   де N — кількість нейронів, P — кількість образів, x_k — k-ий біполярний образ.")
    
    add_paragraph(doc, "3. Асинхронне оновлення:", italic=True)
    add_paragraph(doc, "   s_i(t+1) = sign( Σ(j=1 to N) W_ij * s_j(t) )")
    
    add_paragraph(doc, "4. Функція енергії (функція Ляпунова):", italic=True)
    add_paragraph(doc, "   E(s) = -0.5 * s^T * W * s")
    
    doc.add_page_break()
    
    # Виконання роботи
    add_heading(doc, "ВИКОНАННЯ РОБОТИ", level=1)
    
    add_heading(doc, "Інструменти", level=2)
    add_paragraph(doc, "• Мова програмування: Python 3.11")
    add_paragraph(doc, "• Бібліотеки: NumPy, Matplotlib, Tkinter")
    
    add_heading(doc, "Основні результати", level=2)
    
    add_heading(doc, "1. Варіант та навчальні літери", level=3)
    add_paragraph(doc, "Мій варіант — 1, навчальні літери: B, D, Y.")
    
    add_heading(doc, "2. Шумовий експеримент", level=3)
    add_paragraph(doc, "Проведено експеримент з силою шуму σ = 0.4, 50 випробувань на літеру:")
    
    # Таблиця результатів
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Заголовки
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Літера"
    hdr_cells[1].text = "Точність"
    hdr_cells[2].text = "Збіжність"
    
    # Дані
    row1 = table.rows[1].cells
    row1[0].text = "B"
    row1[1].text = "0.00"
    row1[2].text = "1.00"
    
    row2 = table.rows[2].cells
    row2[0].text = "D"
    row2[1].text = "0.18"
    row2[2].text = "1.00"
    
    row3 = table.rows[3].cells
    row3[0].text = "Y"
    row3[1].text = "0.66"
    row3[2].text = "1.00"
    
    add_paragraph(doc, "Висновки: Літера Y має найкращу точність (66%), літера B — найгіршу (0%).")
    
    add_heading(doc, "3. Триптихи та графіки енергії", level=3)
    add_paragraph(doc, "Триптихи показують три стани: еталон, зашумлений вхід та відновлений образ.")
    add_paragraph(doc, "Графіки енергії демонструють монотонне спадання функції енергії під час асинхронної релаксації.")
    
    add_heading(doc, "4. «Чужа» літера", level=3)
    add_paragraph(doc, "Використовували літеру V (не входила до навчальної вибірки).")
    add_paragraph(doc, "Результат: після релаксації літера V збіглася до найближчого еталону D.")
    
    doc.add_page_break()
    
    # Висновки
    add_heading(doc, "ВИСНОВКИ", level=1)
    add_paragraph(doc, "1. Успішно реалізовано дискретну мережу Хопфілда з асинхронним оновленням.")
    add_paragraph(doc, "2. Мережа здатна відновлювати літери з зашумлених входів, але точність залежить від схожості навчальних образів.")
    add_paragraph(doc, "3. Функція енергії монотонно спадає, що підтверджує теоретичні положення про збіжність мережі.")
    add_paragraph(doc, "4. «Чужі» літери (не з навчальної вибірки) мають тенденцію сходитися до найближчого навчального еталону.")
    add_paragraph(doc, "5. Мережа Хопфілда може бути використана для задач асоціативної пам'яті та розпізнавання простих образів.")
    
    # Save document
    output_path = Path("ЗВІТ_ЛР6_Петренко_Ярослав_КНД-31.docx")
    doc.save(output_path)
    print(f"Звіт студента створено: {output_path.absolute()}")


if __name__ == "__main__":
    create_student_report()
